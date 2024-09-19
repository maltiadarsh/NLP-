import pandas as pd
from bs4 import BeautifulSoup
import requests
import nltk
import os
from docx import Document
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from textblob import TextBlob

# Ensure required NLTK data packages are downloaded
nltk.download('punkt')
nltk.download('stopwords')

# Load the input Excel file
input_file = 'Input.xlsx'
df = pd.read_excel(input_file)

# Define output directories
text_output_dir = 'extracted_articles'
os.makedirs(text_output_dir, exist_ok=True)
doc_output_file = 'Text Analysis.docx'

# Create a new Document for storing the extracted texts
doc = Document()

def extract_text(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # You will need to customize this section to extract the article's title and text
    title = soup.find('title').get_text() if soup.find('title') else 'No Title'
    paragraphs = soup.find_all('p')
    text = ' '.join([para.get_text() for para in paragraphs])
    
    return title, text

for index, row in df.iterrows():
    url_id = row['URL_ID']
    url = row['URL']

    try:
        # Extract the article title and text
        title, text = extract_text(url)

        # Save the text to a text file named after the URL_ID
        with open(os.path.join(text_output_dir, f'{url_id}.txt'), 'w', encoding='utf-8') as file:
            file.write(f'{title}\n{text}')

        # Add the extracted text to the Word document
        doc.add_heading(f'Article {url_id}: {title}', level=1)
        doc.add_paragraph(text)

        print(f'Extracted and saved article {url_id}')
    except Exception as e:
        print(f'Error processing {url_id}: {e}')


# Negative score function definition
def get_sentiment(text):
    analysis = TextBlob(text)
    return analysis.sentiment.polarity

def analyze_text(text):
    # Tokenize into sentences and words
    sentences = sent_tokenize(text)
    words = word_tokenize(text.lower())
    
    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    filtered_words = [word for word in words if word.isalnum() and word not in stop_words]

    # Calculate required variables
    polarity_score = get_sentiment(text)
    positive_score = max(polarity_score, 0)
    negative_score = min(polarity_score, 0)
    subjectivity_score = TextBlob(text).sentiment.subjectivity
    avg_sentence_length = len(filtered_words) / len(sentences) if sentences else 0
    complex_word_count = sum(1 for word in filtered_words if sum(1 for char in word if char in 'aeiou') >= 3)
    percentage_complex_words = (complex_word_count / len(filtered_words)) * 100 if filtered_words else 0
    fog_index = 0.4 * (avg_sentence_length + percentage_complex_words)
    word_count = len(filtered_words)
    syllable_per_word = sum(sum(1 for char in word if char in 'aeiou') for word in filtered_words) / word_count if word_count else 0
    personal_pronouns = len([word for word in words if word in ['i', 'we', 'my', 'ours', 'us']])
    avg_word_length = sum(len(word) for word in filtered_words) / word_count if word_count else 0
    avg_number_of_word_per_sentence = word_count / len(sentences) if sentences else 0

    return {
        'positive_score': positive_score,
        'negative_score': negative_score,
        'polarity_score': polarity_score,
        'subjectivity_score': subjectivity_score,
        'avg_number_of_word_per_sentence': avg_number_of_word_per_sentence,
        'percentage_complex_words': percentage_complex_words,
        'fog_index': fog_index,
        'complex_word_count': complex_word_count,
        'word_count': word_count,
        'syllable_per_word': syllable_per_word,
        'personal_pronouns': personal_pronouns,
        'avg_word_length': avg_word_length
    }

# Perform analysis on each text file and store results
results = []
for filename in os.listdir(text_output_dir):
    if filename.endswith('.txt'):
        url_id = os.path.splitext(filename)[0]
        with open(os.path.join(text_output_dir, filename), 'r', encoding='utf-8') as file:
            text = file.read()
            analysis = analyze_text(text)
            results.append((url_id, analysis))

# Save the document
doc.save(doc_output_file)

# Output to Excel
data_to_append = []

for data in results:
    url_id = data[0]
    analysis = data[1]
    data_to_append.append({
        'URL_ID': url_id,
        'POSITIVE SCORE': analysis['positive_score'],
        'NEGATIVE SCORE': analysis['negative_score'],
        'POLARITY SCORE': analysis['polarity_score'],
        'SUBJECTIVITY SCORE': analysis['subjectivity_score'],
        'AVG SENTENCE LENGTH': analysis['avg_number_of_word_per_sentence'],  
        'PERCENTAGE OF COMPLEX WORDS': analysis['percentage_complex_words'],
        'FOG INDEX': analysis['fog_index'],
        'AVG NUMBER OF WORDS PER SENTENCE': analysis['avg_number_of_word_per_sentence'],
        'COMPLEX WORD COUNT': analysis['complex_word_count'],
        'WORD COUNT': analysis['word_count'],
        'SYLLABLE PER WORD': analysis['syllable_per_word'],
        'PERSONAL PRONOUNS': analysis['personal_pronouns'],
        'AVG WORD LENGTH': analysis['avg_word_length']
    })

# Convert the list of dictionaries to a DataFrame
data_df = pd.DataFrame(data_to_append)

# Save the output DataFrame to an Excel file
data_df.to_excel('Output Data.xlsx', index=False)
print("Mission completed ")
